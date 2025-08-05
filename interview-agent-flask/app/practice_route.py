import json
from flask import Blueprint, request, Response, jsonify, stream_with_context
import PyPDF2
import time
import docx
from io import BytesIO
from services.DeepSeek import DeepseekAPI
from services.SparkPractice import AIPracticeAPI
import os
import datetime


practice_bp = Blueprint('practice', __name__)

@practice_bp.route('/answer', methods=['GET'])
def handle_answer():
    user_message = request.args.get('page', default=1, type=str)
    # user_message = request.json.get('message', '')
    if not user_message:
        return jsonify({'error': 'No message provided'}), 400
    
    # 模拟流式输出
    def generate_response():
        responses = [
            '正在分析您的问题...\n',
            '根据您的描述，我认为...\n',
            '以下是我的建议：\n',
            '1. 首先...\n',
            '2. 其次...\n',
            '3. 最后...\n'
        ]
        for resp in responses:
            time.sleep(0.5)  # 模拟处理延迟
            yield resp.encode('utf-8')
    
    return Response(generate_response(), mimetype='text/event-stream')


@practice_bp.route('/answer_v1', methods=['GET'])
def handle_answer_v1():
    user_message = request.args.get('prompt', default="", type=str)
    if not user_message:
        return jsonify({'error': 'No message provided'}), 400
    
    if user_message:
        # 在用户消息中添加格式要求
        formatted_message = f"{user_message}\n\n请用纯文本格式回复，不要使用Markdown标记符号（如*、#、-等）。"
        
        AIPractice = AIPracticeAPI.getInstance()
        res = AIPractice.get_answer(formatted_message)
        print(res)
        if res:
            return jsonify({'content': res})
        else:
            return jsonify({'error': 'Deepseek API error'}), 500


@practice_bp.route('/answer_v2', methods=['POST'])
def handle_project_packaging_ai():
    """专门用于项目包装的AI分析接口"""
    data = request.get_json()
    
    if not data:
        return jsonify({'error': 'No data provided'}), 400
    
    # 支持两种参数格式
    project_name = data.get('project_name', '') or data.get('projectName', '')
    project_desc = data.get('project_desc', '') or data.get('projectDesc', '')
    file_content = data.get('file_content', '') or data.get('fileContent', '')
    
    if not file_content:
        return jsonify({'error': 'No file content provided'}), 400
    
    # 构建专门的项目包装提示词
    prompt = f'''请你作为一名资深职业规划师，帮助学生将学校项目包装成有竞争力的简历项目。

项目信息：
项目名称：{project_name or '未提供'}
项目描述：{project_desc or '未提供'}

项目文件内容：
{file_content}

请按照以下要求生成项目包装文案：

1. 项目标题：给出一个有吸引力的项目名称
2. 技术栈：列出项目使用的技术和工具
3. 核心功能：总结项目的主要功能模块
4. 技术亮点：突出项目的技术难点和创新点
5. 个人贡献：描述在项目中的具体贡献和角色
6. 项目成果：量化项目的影响和成果

要求：
- 语言专业，突出商业价值
- 使用量化数据（如可能）
- 强调技术深度和复杂度
- 适合写在简历上
- 控制在300字以内
- 请用纯文本格式回复，不要使用Markdown标记符号（如*、#、-等）

请直接输出包装后的项目描述，不要包含其他说明文字。'''

    try:
        # 调用DeepSeek API
        deepseek = DeepseekAPI.getInstance()
        response = deepseek.safe_generate_content_deepseek2(prompt)
        
        if response and response.text:
            return jsonify({
                'success': True,
                'content': response.text
            })
        else:
            return jsonify({
                'success': False,
                'error': 'AI生成内容为空'
            }), 500
            
    except Exception as e:
        error_msg = str(e)
        print(f"项目包装AI分析错误: {error_msg}")
        
        return jsonify({
            'success': False,
            'error': f'AI服务暂时不可用: {error_msg}'
        }), 500


@practice_bp.route('/offer_compare', methods=['POST'])
def offer_compare():
    """Offer对比分析接口"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        offer_a = data.get('offer_a', '')
        offer_b = data.get('offer_b', '')
        priority = data.get('priority', '')
        
        if not offer_a or not offer_b:
            return jsonify({'error': 'Both offers are required'}), 400
        
        # 构建Offer对比分析提示词
        prompt = f'''你是一位资深的职业规划专家和HR顾问，请帮助用户对比分析两个工作机会。

Offer A详情：
{offer_a}

Offer B详情：
{offer_b}

用户关注点：
{priority if priority else '未特别说明'}

请从以下维度进行深度对比分析：

1. 薪资待遇对比（基本工资、奖金、股权、福利等）
2. 发展前景分析（晋升空间、学习机会、行业发展等）
3. 工作环境评估（团队氛围、工作强度、地理位置等）
4. 公司实力对比（规模、稳定性、品牌影响力等）
5. 个人匹配度（技能要求、兴趣匹配、长期规划等）

最后给出明确的建议：推荐选择哪个offer，并说明3个主要理由。

要求：
- 分析要客观全面，考虑短期和长期利益
- 结合用户的关注点进行重点分析
- 语言简洁易懂，避免过多专业术语
- 用纯文本格式回复，不要使用Markdown标记符号
- 总字数控制在500字以内'''

        try:
            # 调用DeepSeek API
            deepseek = DeepseekAPI.getInstance()
            response = deepseek.safe_generate_content_deepseek2(prompt)
            
            if response and response.text:
                return jsonify({
                    'success': True,
                    'content': response.text
                })
            else:
                return jsonify({
                    'success': False,
                    'error': 'AI分析内容为空'
                }), 500
                
        except Exception as e:
            error_msg = str(e)
            print(f"Offer对比AI分析错误: {error_msg}")
            
            return jsonify({
                'success': False,
                'error': f'AI服务暂时不可用: {error_msg}'
            }), 500
            
    except Exception as e:
        return jsonify({'error': f'Request processing failed: {str(e)}'}), 500


@practice_bp.route('/evaluate', methods=['GET'])
def evaluate():
    history_data = request.args.get('historyData', default="", type=str)
    if not history_data:
        return jsonify({'error': 'No history_data provided'}), 400
    
    print(history_data)
    prompt = '''请根据我的根据历史刷题记录分析我的薄弱环节、高频错误点，并给出针对性的提升策略（如重点练习哪些题型、时间管理建议等）。要求分析简洁清晰，建议可操作性强。请用纯文本格式回复，不要使用Markdown标记符号（如*、#、-等）。
                以下是我的历史刷题记录：
                <{}>
             '''.format(history_data)
    print(prompt)
    if history_data:
        Deepseek = DeepseekAPI.getInstance()
        res = Deepseek.safe_generate_content_deepseek2(prompt)
        print(res)
        print(res.text)
        if res:
            return jsonify({'content': res.text})
        else:
            return jsonify({'error': 'Deepseek API error'}), 500

@practice_bp.route('/evaluate_v2', methods=['GET'])
def evaluate_stream():
    history_data = request.args.get('historyData', default="", type=str)
    if not history_data:
        return jsonify({'error': 'No history_data provided'}), 400
    
    print(history_data)
    prompt = '''请根据我的根据历史刷题记录分析我的薄弱环节、高频错误点，并给出针对性的提升策略（如重点练习哪些题型、时间管理建议等）。要求分析简洁清晰，建议可操作性强。请用纯文本格式回复，不要使用Markdown标记符号（如*、#、-等）。
                以下是我的历史刷题记录：
                <{}>
             '''.format(history_data)
    def generate():
        stream = DeepseekAPI.getInstance().global_deepseek_client.chat.completions.create(
            model="deepseek-chat",
            messages=[{"role": "user", "content": prompt}],
            stream=True
        )
        all_text = ""
        for chunk in stream:
            if chunk.choices and chunk.choices[0].delta.content:
                all_text += chunk.choices[0].delta.content
                print(all_text)
                yield f"data: {chunk.choices[0].delta.content}\n\n"
    return Response(generate(), mimetype='text/event-stream', headers = {'Cache-Control': 'no-cache','Connection': 'keep-alive'})


@practice_bp.route('/recommend_questions', methods=['GET'])
def recommend_questions():
    """
    刷题推荐API
    根据用户的专业、目标岗位、难度偏好推荐合适的面试题目
    """
    # 获取请求参数
    major = request.args.get('major', default="计算机科学", type=str)
    position = request.args.get('position', default="软件工程师", type=str) 
    difficulty = request.args.get('difficulty', default="中等", type=str)
    question_type = request.args.get('type', default="技术", type=str)
    
    # 构建推荐提示词
    prompt = f'''你是一位专业的面试官和技术导师。请为{major}专业、目标职位为{position}的学生推荐5道{difficulty}难度的{question_type}类面试题。

要求：
1. 题目内容要具体完整，不要只给标题
2. 每道题目包含：题目描述、考察要点、答题思路
3. 难度要循序渐进，从基础到进阶
4. 要结合用户的专业背景和求职方向
5. 用纯文本格式输出，可以使用数字编号和简单分段

格式示例：
题目1：请介绍一下你对{question_type}的理解
考察要点：基础概念理解
答题思路：从定义、特点、应用场景三个角度回答

题目2：如何优化{position}工作中的性能问题
考察要点：实际应用能力
答题思路：分析瓶颈、提出解决方案、举例说明

请直接生成5道具体的面试题目，让用户可以立即开始练习。每道题目都要包含完整的题目描述，不要使用过多的符号标记。'''

    try:
        # 调用DeepSeek API
        deepseek = DeepseekAPI.getInstance()
        response = deepseek.safe_generate_content_deepseek2(prompt)
        
        if response and response.text:
            # 直接返回AI生成的文本内容，不再尝试解析JSON
            return jsonify({
                'success': True,
                'data': {
                    'questions': [],
                    'raw_content': response.text,
                    'study_suggestions': response.text
                }
            })
        else:
            return jsonify({'success': False, 'error': 'AI生成内容为空'}), 500
            
    except Exception as e:
        error_msg = str(e)
        print(f"刷题推荐API错误: {error_msg}")
        
        # 根据不同错误类型返回不同的错误信息
        if "Connection" in error_msg or "10054" in error_msg:
            return jsonify({
                'success': False, 
                'error': '网络连接不稳定，请稍后重试',
                'error_type': 'connection_error'
            }), 503
        elif "timeout" in error_msg.lower():
            return jsonify({
                'success': False, 
                'error': 'AI服务响应超时，请稍后重试',
                'error_type': 'timeout_error'
            }), 408
        else:
            return jsonify({
                'success': False, 
                'error': f'AI服务暂时不可用: {error_msg}',
                'error_type': 'service_error'
            }), 500


@practice_bp.route('/resume', methods=['POST'])
def handle_resume():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    prompt = request.form.get('prompt', '')
    prompt = '''你是一名资深职业规划师，请根据用户提供的简历内容，按以下要求输出优化建议：
        输出规则：
        1. 格式要求：使用纯文本格式，不要使用Markdown标记符号（如*、#、-等）
        2. 长度控制：每条建议不超过3句话，总输出不超过400字
        3. 内容分级：按优先级标注（关键项 / 改进项 / 加分项）
        4. 禁止事项：不得出现"建议优化"等模糊表述，必须给出具体修改方案
        结尾不要出现"字数统计等字眼"
        '''
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file:
        content = ''
        if file.filename.endswith('.pdf'):
            pdf_reader = PyPDF2.PdfReader(BytesIO(file.read()))
            for page in pdf_reader.pages:
                content += page.extract_text() + '\n'
        elif file.filename.endswith('.docx'):
            doc = docx.Document(file)
            content = '\n'.join([para.text for para in doc.paragraphs])
        # 保存简历内容到本地文件
        save_dir = os.path.join(os.path.dirname(__file__), '../resource/resume')
        timestamp = int(time.time())
        filename = f"resume-{timestamp}.txt"
        filepath = os.path.join(save_dir, filename)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        
        print(content)
        print(prompt)
        if content:
            Deepseek = DeepseekAPI.getInstance()
            res = Deepseek.safe_generate_content_deepseek2(prompt+"，以下是简历内容："+content)
            print(res)
            print(res.text)
            if res:
                return jsonify({'content': res.text})
            else:
                return jsonify({'error': 'Deepseek API error'}), 500
    
    return jsonify({'error': 'Invalid file type. Only PDF files are allowed.'}), 400


@practice_bp.route('/project_packaging', methods=['POST'])
def handle_project_packaging():
    """处理项目包装请求"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    project_name = request.form.get('projectName', '')
    project_desc = request.form.get('projectDesc', '')
    
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file:
        content = ''
        try:
            # 根据文件类型处理文件内容
            if file.filename.endswith('.pdf'):
                pdf_reader = PyPDF2.PdfReader(BytesIO(file.read()))
                for page in pdf_reader.pages:
                    content += page.extract_text() + '\n'
            elif file.filename.endswith(('.docx', '.doc')):
                doc = docx.Document(file)
                content = '\n'.join([para.text for para in doc.paragraphs])
            elif file.filename.endswith(('.txt', '.md')):
                content = file.read().decode('utf-8')
            elif file.filename.endswith(('.js', '.py', '.java', '.cpp', '.html', '.css')):
                content = file.read().decode('utf-8')
            else:
                # 对于其他文件类型，尝试作为文本读取
                try:
                    content = file.read().decode('utf-8')
                except:
                    return jsonify({'error': 'Unsupported file type or encoding'}), 400
            
            # 限制内容长度，避免API调用过大
            if len(content) > 10000:
                content = content[:10000] + "...(内容已截断)"
            
            # 保存项目文件到本地（可选）
            save_dir = os.path.join(os.path.dirname(__file__), '../resource/resume')
            os.makedirs(save_dir, exist_ok=True)
            timestamp = int(time.time())
            filename = f"project-{timestamp}.txt"
            filepath = os.path.join(save_dir, filename)
            
            # 保存项目信息和文件内容
            project_info = f"项目名称：{project_name}\n项目描述：{project_desc}\n\n文件内容：\n{content}"
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(project_info)
            
            print(f"项目文件已保存：{filepath}")
            print(f"项目名称：{project_name}")
            print(f"项目描述：{project_desc}")
            print(f"文件内容长度：{len(content)}")
            
            # 返回文件内容给前端，让前端调用AI分析
            return jsonify({
                'content': content,
                'projectName': project_name,
                'projectDesc': project_desc,
                'message': 'File processed successfully'
            })
            
        except Exception as e:
            print(f"处理文件时出错：{str(e)}")
            return jsonify({'error': f'Error processing file: {str(e)}'}), 500
    
    return jsonify({'error': 'Invalid file'}), 400
